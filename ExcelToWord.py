# -*- coding: utf-8 -*-
"""
Created on Sun Jul 28 15:22:10 2024

@author: lucio.mello
"""
import pandas as pd
import openpyxl 
import docx


tabela_teste = pd.read_excel('C:\sistemas\ExcelToWord\PastaTeste.xlsx', 'Planilha1')
doc = docx.Document('./PastaTeste.docx')


for i in range(0, tabela_teste['Titulo'].count()):
    coluna_titulo = tabela_teste.at[i, 'Titulo']
    coluna_valor = tabela_teste.at[i, 'Valor']
    for paragrafo in doc.paragraphs:
        paragrafo.text = paragrafo.text.replace(coluna_titulo, str(coluna_valor))
        

tab = doc.add_table(tabela_teste.shape[0]+1, tabela_teste.shape[1])
    
for j in range(tabela_teste.shape[-1]):
    tab.cell(0,j).text = tabela_teste.columns[j]


for i in range(tabela_teste.shape[0]):
    for j in range(tabela_teste.shape[-1]):
        tab.cell(i+1,j).text = str(tabela_teste.values[i,j])    
    
doc.save('./PastaTeste.docx')
        
